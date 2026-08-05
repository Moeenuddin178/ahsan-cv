from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_cv():
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header - Name
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run("RAJA AHSAN ALI")
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Position
    position = doc.add_paragraph()
    position.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pos_run = position.add_run("Security Guard")
    pos_run.font.size = Pt(14)
    pos_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Contact info line
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run("Email: ehxanali365@gmail.com  |  Phone: +974 33501566  |  Location: Alsaad Rabiea C, Qatar")
    contact_run.font.size = Pt(10)
    
    # Add line
    doc.add_paragraph("_" * 80)
    
    # Professional Summary
    summary_heading = doc.add_paragraph()
    summary_run = summary_heading.add_run("PROFESSIONAL SUMMARY")
    summary_run.bold = True
    summary_run.font.size = Pt(12)
    summary_run.font.color.rgb = RGBColor(0, 51, 102)
    
    summary_text = doc.add_paragraph()
    summary_text.add_run(
        "Motivated and reliable professional with six months of hands-on experience in Qatar. "
        "Previously worked as a Picker at Snoomart and as a Coordinator at IGI Vitality Insurance Department. "
        "Proven ability to work efficiently in fast-paced environments with strong attention to detail "
        "and excellent coordination skills. Seeking a position as a Security Guard to contribute to "
        "maintaining safety and security standards."
    )
    summary_text.paragraph_format.space_after = Pt(12)
    
    # Personal Information
    personal_heading = doc.add_paragraph()
    personal_run = personal_heading.add_run("PERSONAL INFORMATION")
    personal_run.bold = True
    personal_run.font.size = Pt(12)
    personal_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Personal info table
    personal_table = doc.add_table(rows=3, cols=4)
    personal_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    personal_data = [
        ("Full Name:", "Raja Ahsan Ali", "Age:", "25 years"),
        ("Height:", "5'10\"", "Weight:", "62 kg"),
        ("QID:", "33501566", "Location:", "Alsaad Rabiea C, Qatar")
    ]
    
    for i, row_data in enumerate(personal_data):
        row = personal_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0:  # Labels
                cell.paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # Work Experience
    exp_heading = doc.add_paragraph()
    exp_run = exp_heading.add_run("WORK EXPERIENCE")
    exp_run.bold = True
    exp_run.font.size = Pt(12)
    exp_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Job 1 - Most Recent (IGI)
    job1_title = doc.add_paragraph()
    job1_run = job1_title.add_run("Coordinator")
    job1_run.bold = True
    job1_run.font.size = Pt(11)
    
    job1_company = doc.add_paragraph()
    job1_company.add_run("IGI Vitality Insurance Department | Qatar | Previous Position")
    job1_company.paragraph_format.space_after = Pt(6)
    
    job1_duties = [
        "Coordinated departmental activities and ensured smooth operations",
        "Handled administrative tasks and documentation",
        "Communicated with clients and team members effectively",
        "Managed workflow and maintained organized records"
    ]
    
    for duty in job1_duties:
        p = doc.add_paragraph(duty, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # Job 2 - First Job (Snoomart)
    job2_title = doc.add_paragraph()
    job2_run = job2_title.add_run("Picker")
    job2_run.bold = True
    job2_run.font.size = Pt(11)
    
    job2_company = doc.add_paragraph()
    job2_company.add_run("Snoomart | Qatar | Previous Position")
    job2_company.paragraph_format.space_after = Pt(6)
    
    job2_duties = [
        "Efficiently picked and prepared orders for dispatch",
        "Maintained accuracy in order fulfillment",
        "Worked in a fast-paced warehouse environment",
        "Collaborated with team members to meet daily targets"
    ]
    
    for duty in job2_duties:
        p = doc.add_paragraph(duty, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # Education
    edu_heading = doc.add_paragraph()
    edu_run = edu_heading.add_run("EDUCATION")
    edu_run.bold = True
    edu_run.font.size = Pt(12)
    edu_run.font.color.rgb = RGBColor(0, 51, 102)
    
    edu_text = doc.add_paragraph()
    edu_title = edu_text.add_run("Intermediate (I.Com)")
    edu_title.bold = True
    edu_text.add_run("\nCommerce stream with focus on business and accounting fundamentals")
    
    # Skills
    skills_heading = doc.add_paragraph()
    skills_run = skills_heading.add_run("SKILLS")
    skills_run.bold = True
    skills_run.font.size = Pt(12)
    skills_run.font.color.rgb = RGBColor(0, 51, 102)
    
    skills = [
        "Strong physical fitness and stamina",
        "Excellent observation and attention to detail",
        "Good communication skills (verbal and written)",
        "Ability to work in shifts and under pressure",
        "Team player with coordination experience",
        "Punctual and reliable",
        "Basic computer literacy"
    ]
    
    # Create 2-column layout for skills
    skills_table = doc.add_table(rows=4, cols=2)
    for i, skill in enumerate(skills):
        row_idx = i // 2
        col_idx = i % 2
        if row_idx < 4:
            skills_table.rows[row_idx].cells[col_idx].text = "• " + skill
    
    # Languages
    lang_heading = doc.add_paragraph()
    lang_run = lang_heading.add_run("LANGUAGES")
    lang_run.bold = True
    lang_run.font.size = Pt(12)
    lang_run.font.color.rgb = RGBColor(0, 51, 102)
    
    languages = doc.add_paragraph()
    languages.add_run("• Urdu (Native)    • English (Working proficiency)    • Arabic (Basic)")
    
    # References
    ref_heading = doc.add_paragraph()
    ref_run = ref_heading.add_run("REFERENCES")
    ref_run.bold = True
    ref_run.font.size = Pt(12)
    ref_run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph("Available upon request")
    
    # Save document
    doc.save('docs/Raja_Ahsan_Ali_CV.docx')
    print("CV created successfully: docs/Raja_Ahsan_Ali_CV.docx")

if __name__ == "__main__":
    create_cv()
